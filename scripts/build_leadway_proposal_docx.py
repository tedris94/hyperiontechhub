"""One-off generator: Leadway proposal -> DOCX."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def main():
    out_dir = Path(__file__).resolve().parent.parent / "proposals"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Leadway-College-Platform-Proposal-Hyperion-Tech-Hub.docx"

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Technical & Financial Proposal", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Website & Integrated Platform for Programmes, Registration, Admissions, Examinations & Payments")
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        "Prepared for: Leadway College of Advanced Studies Limited",
        "(formerly Leadway Educational Consults)",
        "Prepared by: Hyperion Tech Hub",
        "Date: 10 May 2026",
    ):
        meta.add_run(line + "\n")

    doc.add_page_break()

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "Hyperion Tech Hub proposes an end-to-end digital platform aligned with your focus on "
        "IJMB, JAMB, WAEC, NECO lessons and related services. The solution replaces ad-hoc forms and "
        "manual tracking with a single system: a modern public website, secure student/parent "
        "portals, staff dashboards, structured admissions and exams, and online/offline-capable "
        "payments suitable for the Nigerian market."
    )
    doc.add_paragraph(
        "References: https://leadwayeducation.wixsite.com/leadway/ ; "
        "https://web.facebook.com/p/Leadway-Education-Consult-61555663662917/"
    )

    doc.add_heading("2. Understanding of your business", level=1)
    doc.add_paragraph(
        "From public information, Leadway offers tutoring and exam preparation, emphasises "
        "personalised learning, and serves families in the Abuja/FCT area with multiple contact "
        "channels. The proposed system supports that positioning while scaling operations (more "
        "cohorts, branches, or online students) without proportional administrative overhead."
    )

    doc.add_heading("3. Objectives", level=1)
    objectives = [
        ("Professional brand presence", "Fast, credible website with clear programme pages, FAQs, and lead capture."),
        ("Programme management", "Define tracks (e.g. IJMB, JAMB, WAEC, NECO), batches, timetables, and fees in one place."),
        ("Registration & admissions", "Online applications, document upload, status workflow, and communication."),
        ("Examinations & academics", "Schedules, attendance (where applicable), mock tests, results, and basic reporting."),
        ("Payments", "Invoicing, part-payments, receipts, reconciliation, and integration with major Nigerian gateways."),
        ("Compliance & trust", "Role-based access, audit logs, backups, and privacy-conscious handling of student data."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Objective"
    hdr[1].text = "Outcome"
    for obj, out in objectives:
        row = table.add_row().cells
        row[0].text = obj
        row[1].text = out

    doc.add_heading("4. Proposed solution architecture (high level)", level=1)
    doc.add_heading("4.1 Channels", level=2)
    doc.add_paragraph(
        "Marketing website: SEO-friendly pages, programme detail, testimonials, blog/news, contact, "
        "and Apply/Register entry points.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Web application (responsive): Full functionality for applicants, students, parents/guardians, "
        "tutors, and administrators on modern browsers.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Mobile: Progressive Web App (PWA) for lower cost and single codebase, installable on phones; "
        "or native iOS/Android apps for richer offline/push features. Phase 1 typically starts with PWA.",
        style="List Bullet",
    )

    doc.add_heading("4.2 Core modules", level=2)
    modules = [
        "Identity & access — Accounts, roles (super admin, registrar, finance, tutor, student, parent), optional two-factor for staff.",
        "Programmes & catalog — Programmes, levels, subjects, fee structures, discounts, instalments.",
        "Registration & admissions — Multi-step forms, file uploads (credentials), application fees, pipeline statuses (submitted → review → offer → enrolled).",
        "Enrolment & classes — Cohorts, batches, timetable, room/online links, tutor assignment.",
        "Exams & assessment — Exam calendar, question banks (optional), online quizzes or paper-based result entry, transcripts/report cards (scope-dependent).",
        "Payments — Paystack/Flutterwave (or agreed provider), webhooks for automatic confirmation, invoices, refunds workflow (policy-driven).",
        "Notifications — Email; optional SMS/WhatsApp Business API for reminders and payment alerts.",
        "Reporting — Revenue by programme, outstanding fees, admission funnel, exam performance summaries.",
    ]
    for m in modules:
        doc.add_paragraph(m, style="List Number")

    doc.add_heading("4.3 Technical approach (illustrative)", level=2)
    doc.add_paragraph(
        "Frontend: React or Next.js. Backend: Node.js or Laravel (final choice follows team expertise). "
        "Database: PostgreSQL or MySQL with regular encrypted backups. Hosting: managed cloud with staging "
        "and production. Security: HTTPS, OWASP-aligned practices, least-privilege access, encrypted sensitive "
        "fields where appropriate. Exact stack confirmed in a short discovery workshop."
    )

    doc.add_heading("5. Implementation phases & timeline (indicative)", level=1)
    phases = [
        ("Discovery & UX", "Workflows, roles, data model, wireframes, brand alignment", "2–3 weeks"),
        ("MVP", "Website + auth + programmes + online registration + payments + admin dashboard", "8–12 weeks"),
        ("Admissions depth", "Document verification workflow, offers, contracts/digital acceptance", "3–5 weeks"),
        ("Exams & academics", "Scheduling, results entry, parent/student views", "4–8 weeks"),
        ("Hardening & launch", "UAT, performance, training, go-live, hypercare", "2–3 weeks"),
    ]
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    h2[0].text = "Phase"
    h2[1].text = "Scope"
    h2[2].text = "Duration (indicative)"
    for a, b, c in phases:
        row = t2.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    doc.add_paragraph(
        "Total (typical): approximately 4–6 months from kickoff to full feature set, with an earlier "
        "MVP launch (~3 months) if registration and payments are prioritised first."
    )

    doc.add_heading("6. Deliverables", level=1)
    for item in (
        "Production website and web app (and PWA or native apps per agreement).",
        "Administrator and role-based documentation (user guides).",
        "Source code ownership transfer or escrow per contract.",
        "Deployment runbooks, backup, and monitoring baseline.",
        "Training session(s) for staff.",
        "Warranty period (e.g. 30–90 days) for defect fixes on agreed scope.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Financial proposal (indicative — NGN)", level=1)
    doc.add_paragraph(
        "Figures are order-of-magnitude estimates for a Nigerian education SME; final pricing follows "
        "discovery (user volumes, exam complexity, integrations, and design depth)."
    )
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    h3 = t3.rows[0].cells
    h3[0].text = "Package"
    h3[1].text = "What’s included"
    h3[2].text = "Indicative investment (NGN)"
    packages = [
        (
            "A — Foundation",
            "New website + programme pages + contact/leads + basic admin CMS + 1 payment integration + hosting setup guidance",
            "₦2.5M – ₦5.0M",
        ),
        (
            "B — Operations (recommended)",
            "Foundation + full registration/admissions workflow + student/parent portals + finance dashboard + notifications (email; optional SMS budget separate)",
            "₦6.0M – ₦12.0M",
        ),
        (
            "C — Full academic suite",
            "Package B + exams/scheduling/results + richer reporting + PWA + optional API for future mobile",
            "₦12.0M – ₦22.0M",
        ),
        (
            "D — Native mobile add-on",
            "iOS + Android apps sharing backend with web",
            "+₦4.0M – ₦10.0M (on top of B or C)",
        ),
    ]
    for a, b, c in packages:
        row = t3.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    doc.add_paragraph("Ongoing costs (pass-through / operational):", style="Heading 3")
    doc.add_paragraph(
        "Hosting & infrastructure: roughly ₦50k – ₦250k/month depending on traffic and redundancy.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Third-party fees: payment gateway charges, SMS, email delivery, domain, SSL.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Retainer (optional): ₦150k – ₦600k/month for security updates, small enhancements, and priority support (hours cap to be agreed).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Payment terms (typical): 40% on contract signing, 30% on MVP/UAT milestone, 30% on go-live — "
        "or monthly milestones for longer engagements."
    )

    doc.add_heading("8. Assumptions & dependencies", level=1)
    assumptions = [
        "Client provides timely content (copy, images, fee tables, programme rules), branding assets, and a named product owner.",
        "Payment provider KYC and settlement accounts are completed by the client.",
        "Legal templates (enrolment agreements, refund policy) are supplied or reviewed by client’s counsel; Hyperion implements them as digital flows only.",
        "Exam integrity (e.g. proctored high-stakes online exams) may require specialised third-party tools; scope and cost are separate if required.",
    ]
    for a in assumptions:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("9. Risk & quality", level=1)
    doc.add_paragraph(
        "Phased delivery reduces risk: revenue-related features (registration + payments) go live first. "
        "Staging environment and UAT sign-off before production. Data migration from Wix/manual spreadsheets "
        "can be scoped as a discrete work package if historical records must move over."
    )

    doc.add_heading("10. Next steps", level=1)
    for i, step in enumerate(
        (
            "Discovery call (60–90 minutes) — map current process from first enquiry to enrolment to exams.",
            "Written scope & fixed quote for chosen package (A–C) plus any add-ons.",
            "Contract & kickoff — access, branding, and milestone schedule.",
        ),
        start=1,
    ):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    doc.add_heading("Contact", level=1)
    doc.add_paragraph("Hyperion Tech Hub — [insert email] | [insert phone] | [insert website]")

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
